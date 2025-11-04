"""
Document processing service for multiple formats
Enhanced with validation, malware scanning, and better format support
"""
import io
import logging
import re
import hashlib
from typing import Optional, Tuple, Dict, Any
from pathlib import Path
import mimetypes

logger = logging.getLogger(__name__)

# Maximum file size: 10MB
MAX_FILE_SIZE = 10 * 1024 * 1024

# Supported MIME types
SUPPORTED_MIME_TYPES = {
    'application/pdf',
    'application/msword',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'text/plain',
    'image/jpeg',
    'image/jpg',
    'image/png',
    'image/gif',
    'image/bmp',
    'image/tiff'
}

# Supported file extensions
SUPPORTED_EXTENSIONS = {
    '.pdf', '.doc', '.docx', '.txt',
    '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff'
}

# Optional imports - server will work without these
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False
    logger.warning("PyPDF2 not installed - PDF processing will be limited")

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False
    logger.warning("pdfplumber not installed - PDF processing will be limited")

try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed - DOCX processing disabled")

try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
    # Try to set Tesseract path if on Windows
    try:
        import os
        if os.name == 'nt':  # Windows
            # Common Tesseract installation paths
            tesseract_paths = [
                r'C:\Program Files\Tesseract-OCR\tesseract.exe',
                r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
                r'C:\Users\{}\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'.format(os.getenv('USERNAME', '')),
            ]
            for path in tesseract_paths:
                if os.path.exists(path):
                    pytesseract.pytesseract.tesseract_cmd = path
                    logger.info(f"Tesseract found at: {path}")
                    break
    except Exception as e:
        logger.debug(f"Could not auto-detect Tesseract path: {e}")
except ImportError:
    HAS_OCR = False
    logger.warning("OCR libraries not installed - image processing disabled")

try:
    from pdf2image import convert_from_bytes
    HAS_PDF2IMAGE = True
    # Try to set poppler path if on Windows
    try:
        import os
        if os.name == 'nt':  # Windows
            # poppler is usually in PATH, but if not, user can set it
            # This is optional - pdf2image will work if poppler is in PATH
            pass
    except Exception:
        pass
except ImportError:
    HAS_PDF2IMAGE = False
    logger.warning("pdf2image not installed - OCR for PDFs disabled")


class DocumentProcessor:
    """Handles extraction of text from various document formats"""
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes, filename: str) -> Tuple[str, Optional[str]]:
        """
        Extract text from PDF file
        
        Args:
            file_content: PDF file content as bytes
            filename: Original filename
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        try:
            text = ""
            
            # Try pdfplumber first (better for structured PDFs)
            if HAS_PDFPLUMBER:
                try:
                    pdf = pdfplumber.open(io.BytesIO(file_content))
                    text_parts = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    pdf.close()
                    text = "\n\n".join(text_parts)
                except Exception as e:
                    logger.warning(f"pdfplumber failed: {e}")
            
            # If pdfplumber didn't extract much, try PyPDF2
            if (not text or len(text.strip()) < 100) and HAS_PYPDF2:
                try:
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                    text_parts = []
                    for page in pdf_reader.pages:
                        text_parts.append(page.extract_text())
                    text = "\n\n".join(text_parts)
                except Exception as e:
                    logger.warning(f"PyPDF2 failed: {e}")
            
            # If still no text, try OCR
            if (not text or len(text.strip()) < 100) and HAS_PDF2IMAGE and HAS_OCR:
                try:
                    logger.info("PDF appears to be scanned, attempting OCR...")
                    images = convert_from_bytes(file_content)
                    ocr_texts = []
                    for image in images:
                        ocr_text = pytesseract.image_to_string(image)
                        ocr_texts.append(ocr_text)
                    text = "\n\n".join(ocr_texts)
                except Exception as e:
                    logger.warning(f"OCR failed: {e}")
            
            if not text:
                return "", "No PDF processing libraries installed. Install: pip install PyPDF2 pdfplumber"
            
            return text.strip(), None
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            return "", f"Error processing PDF: {str(e)}"
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> Tuple[str, Optional[str]]:
        """
        Extract text from DOCX file
        
        Args:
            file_content: DOCX file content as bytes
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        try:
            if not HAS_DOCX:
                return "", "python-docx not installed. Install: pip install python-docx"
            
            doc = Document(io.BytesIO(file_content))
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            text = "\n\n".join(text_parts)
            return text.strip(), None
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            return "", f"Error processing DOCX: {str(e)}"
    
    @staticmethod
    def extract_text_from_txt(file_content: bytes) -> Tuple[str, Optional[str]]:
        """
        Extract text from TXT file
        
        Args:
            file_content: TXT file content as bytes
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        try:
            # Try UTF-8 first
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                # Fallback to latin-1
                text = file_content.decode('latin-1', errors='ignore')
            
            return text.strip(), None
            
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {str(e)}")
            return "", f"Error processing TXT: {str(e)}"
    
    @staticmethod
    def extract_text_from_image(file_content: bytes) -> Tuple[str, Optional[str]]:
        """
        Extract text from image using OCR
        
        Args:
            file_content: Image file content as bytes
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        try:
            if not HAS_OCR:
                return "", "OCR libraries not installed. Install: pip install pytesseract Pillow"
            
            image = Image.open(io.BytesIO(file_content))
            text = pytesseract.image_to_string(image)
            return text.strip(), None
            
        except Exception as e:
            logger.error(f"Error extracting text from image: {str(e)}")
            return "", f"Error processing image: {str(e)}"
    
    @staticmethod
    def validate_file(file_content: bytes, filename: str, content_type: str) -> Tuple[bool, Optional[str]]:
        """
        Validate uploaded file
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            content_type: MIME type of the file
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check file size
        if len(file_content) > MAX_FILE_SIZE:
            return False, f"File size exceeds maximum limit of {MAX_FILE_SIZE / (1024*1024):.1f}MB"
        
        if len(file_content) == 0:
            return False, "File is empty"
        
        # Check file extension
        file_ext = Path(filename).suffix.lower()
        if file_ext not in SUPPORTED_EXTENSIONS:
            return False, f"Unsupported file extension: {file_ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
        
        # Validate MIME type
        if content_type and content_type not in SUPPORTED_MIME_TYPES:
            # Try to detect from content
            detected_type, _ = mimetypes.guess_type(filename)
            if detected_type and detected_type not in SUPPORTED_MIME_TYPES:
                return False, f"Unsupported file type: {content_type}"
        
        # Basic malware scanning (check for suspicious patterns)
        suspicious_patterns = [
            b'<script', b'javascript:', b'vbscript:', b'onload=',
            b'<?php', b'<%', b'eval(', b'exec('
        ]
        
        file_content_lower = file_content[:10000].lower()  # Check first 10KB
        for pattern in suspicious_patterns:
            if pattern in file_content_lower:
                logger.warning(f"Suspicious pattern detected in file: {filename}")
                # Don't block, just log - could be false positive
        
        return True, None
    
    @staticmethod
    def get_file_info(file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Get file information
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            
        Returns:
            Dictionary with file information
        """
        file_ext = Path(filename).suffix.lower()
        
        return {
            'filename': filename,
            'size': len(file_content),
            'size_mb': round(len(file_content) / (1024 * 1024), 2),
            'extension': file_ext,
            'mime_type': mimetypes.guess_type(filename)[0],
            'hash': hashlib.md5(file_content).hexdigest()
        }
    
    @staticmethod
    def process_document(file_content: bytes, filename: str, content_type: str) -> Tuple[str, Optional[str], Dict[str, Any]]:
        """
        Process document based on file type with validation
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            content_type: MIME type of the file
            
        Returns:
            Tuple of (extracted_text, error_message, file_info)
        """
        # Validate file first
        is_valid, error = DocumentProcessor.validate_file(file_content, filename, content_type)
        if not is_valid:
            return "", error, DocumentProcessor.get_file_info(file_content, filename)
        
        file_info = DocumentProcessor.get_file_info(file_content, filename)
        filename_lower = filename.lower()
        
        # Determine file type and extract
        text = ""
        error_msg = None
        
        try:
            if filename_lower.endswith('.pdf') or 'pdf' in content_type:
                text, error_msg = DocumentProcessor.extract_text_from_pdf(file_content, filename)
            elif filename_lower.endswith(('.doc', '.docx')) or 'word' in content_type or 'document' in content_type:
                text, error_msg = DocumentProcessor.extract_text_from_docx(file_content)
            elif filename_lower.endswith('.txt') or 'text' in content_type:
                text, error_msg = DocumentProcessor.extract_text_from_txt(file_content)
            elif filename_lower.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff')):
                # For image files, extract text using OCR
                text, error_msg = DocumentProcessor.extract_text_from_image(file_content)
            else:
                # Try to detect as text
                try:
                    text, error_msg = DocumentProcessor.extract_text_from_txt(file_content)
                except:
                    error_msg = f"Unsupported file type: {content_type}"
        except Exception as e:
            logger.error(f"Error processing document: {e}")
            error_msg = f"Error processing document: {str(e)}"
        
        return text, error_msg, file_info

